from telegram import Update, Bot, InlineKeyboardButton, InlineKeyboardMarkup, BotCommand, BotCommandScopeAllGroupChats
from telegram import InputFile
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes
from telegram.ext import ConversationHandler, JobQueue
from sqlalchemy import select, exists
import yaml, requests, time, asyncio, re, traceback, os

import fitz  # PyMuPDF 用來讀PDF
import docx  # python-docx 用來讀Word

from server_main import Individual_search
from Database.MoneyDJ import MoneyDJ
from Database.DB import DB
from utils.AI import GroqAI
from utils.Logger import setup_logger
from DevFeat.news_parser import NewsParser

def shorten_url_tinyurl(long_url):
    api_url = "http://tinyurl.com/api-create.php"
    params = {'url': long_url}
    response = requests.get(api_url, params=params)
    return response.text

def escape_markdown_v2(text: str) -> str:
    escape_chars = r'\_*[]()~`>#+-=|{}.!'
    return ''.join(['\\' + c if c in escape_chars else c for c in text])

def is_valid_input(s):
    return bool(re.fullmatch(r"[1-9][0-9]{3}", s))

def read_pdf(path):
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    return text

def read_word(path):
    doc = docx.Document(path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def clean_markdown(text):
    # 移除連結，只留文字
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # 移除粗體、斜體、底線、刪除線等標記
    text = re.sub(r'(\*|_|~|`)+', '', text)
    return text

class TelegramBot:
    def __init__(self):
        self.TOKEN = yaml.safe_load(open('token.yaml'))["TelegramToken"][0]
        self.group_id = yaml.safe_load(open('token.yaml'))["ChatID"][0]
        self.NP = NewsParser()
        self.lock = asyncio.Lock()
        self.job_queue = JobQueue()
        self.subscribers = set()
        self.db = DB()
        self.groq = GroqAI()
        self.telebot = Bot(token=self.TOKEN)
        self.ASK_CODE = 1
        self.logger = setup_logger()
        # self.report_func = {self.NP.get_uanalyze_report, self.NP.get_fugle_report, self.NP.get_vocus_ieobserve_articles}
        self.report_urls = [
            "https://blog.fugle.tw/",
            "https://feed.cqd.tw/vocus/user/ieobserve", # source: https://github.com/CQD/feeder
            "https://feed.cqd.tw/vocus/user/miula",
            "https://feed.cqd.tw/vocus/user/65ab564cfd897800018a88cc",
            "https://morss.it/:proxy:items=%7C%7C*[class=article-content__title]/https://uanalyze.com.tw/articles",
            "https://morss.it/:proxy/https://www.macromicro.me/blog",
            "https://morss.it/:proxy/https://fintastic.trading/",
        ]
        self.bot_cmd = {"start" : "開始使用機器人", 
                        "help" : "使用說明",
                        "esti" : "估算股票", 
                        "news" : "查看新聞", 
                        "info" : "查詢公司資訊",
                        "subscribe" : "訂閱即時新聞", "unsubscribe" : "取消訂閱即時新聞", "news_summary" : "新聞摘要"}      
        # set source 
        self.news_src_urls = {
                '[udn] 產業' : 'https://morss.it/:proxy:items=%7C%7C*[class=story__headline]/https://money.udn.com/rank/newest/1001/5591/1', 
                '[udn] 證券' : 'https://morss.it/:proxy:items=%7C%7C*[class=story__headline]/https://money.udn.com/rank/newest/1001/5590/1',
                '[udn] 國際' : 'https://morss.it/:proxy:items=%7C%7C*[class=story__headline]/https://money.udn.com/rank/newest/1001/5588/1',
                '[udn] 兩岸' : 'https://morss.it/:proxy:items=%7C%7C*[class=story__headline]/https://money.udn.com/rank/newest/1001/5589/1',
                '[鉅亨網] 頭條' : 'https://api.cnyes.com/media/api/v1/newslist/category/headline',
                '[鉅亨網] 台股' : 'https://api.cnyes.com/media/api/v1/newslist/category/tw_stock',
                '[鉅亨網] 美股' : 'https://api.cnyes.com/media/api/v1/newslist/category/wd_stock',
                '[鉅亨網] 科技' : 'https://api.cnyes.com/media/api/v1/newslist/category/tech',
                "News Digest AI（中文）" : "https://feed.cqd.tw/ndai",
                '[moneydj] 發燒頭條' : 'https://www.moneydj.com/KMDJ/RssCenter.aspx?svc=NR&fno=1&arg=MB010000',
                'Yahoo TW' : "https://tw.stock.yahoo.com/rss?category=news",
            }
        self.news_data = { news_type : [] for news_type in self.news_src_urls.keys() }

    async def set_main_menu(self, application):
        commands = []
        for k, v in self.bot_cmd.items():
            commands.append(BotCommand(k, v))

        await application.bot.set_my_commands(
            commands,
            scope=BotCommandScopeAllGroupChats()
        )

    # 定義 /start 命令處理器
    async def cmd_start(self, update: Update, context):
        # 檢查訊息來源是群組還是私人訊息
        if update.message.chat.type == "group":
            await update.message.reply_text(f"Hello, {update.message.chat.title}! I'm your bot.")
        else:
            await update.message.reply_text('Hello! I am your bot! How can I assist you today?')
    # 定義 /help 命令處理器
    async def cmd_help(self, update: Update, context):
        
        await context.bot.send_message(chat_id=self.group_id, text="help command")
        if update.message.chat.type == "group":
            await update.message.reply_text(f"In this group, I can assist you with commands like /start and /help.")
        else:
            await update.message.reply_text('To use this bot, just type a message, or use /start and /help.')
    # /subscribe
    async def subscribe(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        chat_id = update.effective_chat.id
        self.subscribers.add(chat_id)
        await update.message.reply_text("✅ 已訂閱新聞通知！")
    # /unsubscribe
    async def unsubscribe(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        chat_id = update.effective_chat.id
        self.subscribers.discard(chat_id)
        await update.message.reply_text("❌ 已取消訂閱。")
    # /estimate
    async def cmd_esti(self, update: Update, context):
        # print(context.args)
        stock_list = [x for idx, x in enumerate(context.args) if idx % 2 == 0]
        eps_list = [x for idx, x in enumerate(context.args) if idx % 2 != 0]
        # print(stock_list, eps_list)

        # 檢查訊息來源是群組還是私人訊息
        if update.message.chat.type == "group":
            await update.message.reply_text(f"Hello, {update.message.chat.title}! I'm your bot.")

            for (stock_id, eps) in zip(stock_list, eps_list):
                
                await update.message.reply_text(f"Estimate start: {stock_id}")
                data = Individual_search([stock_id], [eps]) #TODO
                await update.message.reply_text(f"Estimate done: {stock_id}")
        else:
            pass
    # user key info_start
    async def cmd_info(self, update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
        await update.message.reply_text("請輸入股票代碼：")
        return self.ASK_CODE 

    async def cmd_handle_info(self, update: Update, context):
        ticker = update.message.text.strip()
        res = is_valid_input(ticker)
        msg = ""
        
        if ticker is None or res is False:
            msg = "[ERROR] Wrong ticker information"
            await update.message.reply_text(msg)
            return ConversationHandler.END
        
        msg = f"你輸入的代碼是 {ticker}，幫你處理！"
        await update.message.reply_text(msg)
        DJ = MoneyDJ()

        ticker_name, wiki_result = DJ.get_wiki_result(ticker)
        # error handle
        if ticker_name is None or wiki_result is None:
            await update.message.reply_text(f"Information of Ticker {ticker} is not found.")
        else:
            condition = "重點摘要，營收占比或業務占比，有詳細數字的也要列出來"
            prompt = "\n" + condition  + "，並且使用繁體中文回答\n"
            content = self.groq.talk(prompt, wiki_result, reasoning=True)
            # TODO
            save_path = "./files/"
            file_path = f"{save_path}/{str(ticker)}{ticker_name}_info.md"
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            with open(file_path, "rb") as f:
                await update.message.reply_document(
                    document=InputFile(f, filename=file_path),
                    caption="這是你的報告 📄"
                )
        return ConversationHandler.END
    # 定義普通文字訊息處理器
    async def cmd_echo(self, update: Update, context):
        print(f"Received message: {update.message.text}")
        # 群組中的回應
        if update.message.chat.type == "group":
            await update.message.reply_text(f"Group Message: {update.message.text}")
        else:
            await update.message.reply_text(f'You said: {update.message.text}')
    # button callback
    async def button_cb(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        self.logger.debug("✅ Callback handler triggered")
        query = update.callback_query
        
        # response = requests.head(query.data, allow_redirects=True) # 取得最終網址
        
        # article = NP.fetch_news_content(response.url)
        # content = article['content']
        # summary = self.groq.talk(prompt="幫我摘要內容200字以內", content=content, reasoning=True)
        if query.data in self.news_data.keys():
            await query.answer()
            text = ""
            data = self.news_data[query.data]
            for article in data:
                text += f"📰[{escape_markdown_v2(article['title'])}]({article['url']})\n"
            if text != "":
                text = escape_markdown_v2(query.data) + "\n" + text
                await query.edit_message_text(text=text,
                                            parse_mode='MarkdownV2',
                                            disable_web_page_preview=True,
                                            reply_markup=query.message.reply_markup)
        else:
            await query.answer(text="處理中...，以私人回覆方式傳送摘要")
            user = query.from_user
            pass
        # await query.message._bot.send_message(chat_id=user.id, text=f"{article['title']}\n🧠 新聞摘要：\n{summary}")
    # 定義錯誤處理器
    async def cmd_error(self, update: Update, context):
        self.logger.error(context.error)
        
    # 取消對話
    async def cmd_cancel(self, update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
        await update.message.reply_text("已取消操作。")
        return ConversationHandler.END
    # 定義發送新聞的函數
    async def send_news(self, news):

        self.logger.debug("send_news fo subscriber")
        for chat_id in self.subscribers:
            title = news['title']
            url   = news['url']
            titles = f"📰[{escape_markdown_v2(title)}]({url})\n"

            if titles != "":
                text = f"{titles}"
                await self.telebot.send_message(chat_id=chat_id
                                        , text=text
                                        , parse_mode='MarkdownV2')

    async def cmd_news(self, update: Update, context):
        buttons = []
        for key in self.news_data.keys():
            buttons.append(InlineKeyboardButton(key, callback_data=key))

        # keyboard = [[btn] for btn in buttons]
        keyboard = [buttons[i:i+3] for i in range(0, len(buttons), 3)]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message._bot.sendMessage(chat_id=update.message.chat_id
                                            , text="請選擇新聞來源與類型"
                                            , reply_markup=reply_markup
                                            , parse_mode='MarkdownV2')
        
    async def cmd_uanalyze(self, update: Update, context):

        reports = self.NP.get_uanalyze_report()
        for rep  in reports:
            text = f"📰{rep['title']}\n{rep['link']}"
            await update.message._bot.sendMessage(chat_id=update.message.chat_id, text=text)

    async def scheduled_task(self, context: ContextTypes.DEFAULT_TYPE):
        # job_data = context.job.data
        # chat_id = job_data['chat_id']
        # await context.bot.send_message(chat_id=chat_id, text="📢 定時訊息！")
        chatbot = GroqAI()
        prompt = "100字摘要，重要數字也要，且使用繁體中文回答"

        job_data = context.job.data  # 這裡就是你當初設的資料
        chat_id = job_data["chat_id"]
        filenames = []

        for typ, news_list in self.news_data.items():
            filename = typ.replace(" ", "_") + ".md"
            filenames.append(filename)
            with open(filename, "w", encoding="utf-8") as f:
                for news in news_list:
                    f.write(news['title'] + "\n")
                    
                    if 'content' in news.keys():
                        summary = chatbot.talk(prompt=prompt, content=news['content'], reasoning=True)
                        f.write(summary + "\n")
                    else:
                        pass
                    f.write(news['url'] + "\n")
                    f.write('-----------' + "\n")
            await context.bot.send_document(chat_id=chat_id, document=filename)
            time.sleep(5)

    async def cmd_news_summary(self, update, context):
        context.job_queue.run_repeating(self.scheduled_task, interval=60*30, first=0, data={"chat_id": update.effective_chat.id})
        await update.message.reply_text("設定每 30 分鐘傳送摘要檔案！(content available only)")

    async def get_reports(self, context: ContextTypes.DEFAULT_TYPE):


        for idx, url in enumerate(self.report_urls):
            report = self.NP.fetch_report(url)[0]
            exists = self.db.checkReport(report)

            if not exists:
                article = f"{report['title']}\n{report['url']}"
                await context.bot.send_message(chat_id=self.group_id, text=article)
                self.logger.debug("update report")


    async def get_news(self, context: ContextTypes.DEFAULT_TYPE):
        # async with self.lock:
        for news_type, url in self.news_src_urls.items():
            self.logger.info(f"NEWS SOURCE : {news_type}")
            news_list = self.NP.fetch_news_list(url) # Get news from parser only

            titles = [news['title'] for news in self.news_data[news_type]]
            if len(titles) != 0:
                for ele in news_list:
                    if ele['title'] not in titles and not self.subscribers:
                        self.logger.info("SEND NEWS")
                        await self.send_news(ele)
            else:
                self.logger.debug("No news")

            self.news_data[news_type] = news_list # update news

            # update DB
            for news in news_list:
                self.db.checkNews(news)

        self.logger.info("Fetch news sources done")
    
    async def handle_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        if update.message.chat.type == "group":
            return  # 忽略群組中的訊息
        
        # import aspose.words as aw
        if update.message.document:
            document = update.message.document
            file_name = document.file_name.lower()

            # 把檔案下載下來
            file_path = f"./{file_name}"
            file = await document.get_file()   # 第一次 await，拿到檔案物件
            await file.download_to_drive(file_path)  # 第二次 await，下載到本地
            file_name_clear = file_name.split("_", 1)[1]
            await self.telebot.send_message(chat_id=self.group_id, text=f"[TEST]有用戶傳了{file_name_clear}給我，幫你摘要內容")
            # 判斷副檔名
            text = ""
            if file_name.endswith('.pdf'):
                text = read_pdf(file_path)[:8000]
            elif file_name.endswith('.doc') or file_name.endswith('.docx'):
                text = read_word(file_path)[:8000]
            else:
                # await update.message.reply_text("這個檔案格式我還不支援喔！")
                return
            os.remove(file_path)

            summary = self.groq.talk(prompt="幫我做重點摘要500字以內，重點數字優先", content=text, reasoning=True)
            file_path = "./summary.md"
            with open(file_path, "w", encoding="utf-8") as file:
                file.write(summary)
            # doc = aw.Document(file_path)
            # doc.save("summary.pdf")
            with open(file_path, "rb") as file:
                await self.telebot.send_document(chat_id=self.group_id, document=file, caption="這是你的摘要 📄")
            os.remove(file_path)

        elif update.message.photo:
            photo = update.message.photo[-1]
            file = await photo.get_file()
            await context.bot.send_photo(chat_id=self.group_id, photo=file.file_id) # 直接轉傳

        elif update.message.text:
            text = update.message.text
            if "call memo" in text.lower() or "memo" in text.lower():
                await self.telebot.send_message(chat_id=self.group_id, text=f"[TEST]有用戶傳了Call Memo給我，幫你摘要內容")
                summary = self.groq.talk(prompt="幫我做重點摘要500字以內，重點數字優先", content=text, reasoning=True)
                file_path = "./summary.md"
                with open(file_path, "w", encoding="utf-8") as file:
                    file.write(summary)
                with open(file_path, "rb") as file:
                    await self.telebot.send_document(chat_id=self.group_id, document=file, caption="這是你的摘要 📄")
                os.remove(file_path)
            else:
                pass
                # await update.message.reply_text("你傳了一段文字。")
        else:
            pass
            # await update.message.reply_text("這種類型我還看不懂喔。")

    def run(self):

        # 初始化 Application
        application = Application.builder().token(self.TOKEN).build()

        conv_handler = ConversationHandler(
            entry_points=[CommandHandler("info", self.cmd_info)],
            states={
                self.ASK_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, self.cmd_handle_info)]
            },
            fallbacks=[CommandHandler("cancel", self.cmd_cancel)],
        )

        # 註冊處理命令
        application.add_handler(CommandHandler("start", self.cmd_start))
        application.add_handler(CommandHandler("help", self.cmd_help))
        application.add_handler(CommandHandler("esti", self.cmd_esti))
        application.add_handler(CommandHandler("news", self.cmd_news))
        application.add_handler(CommandHandler("uanalyze", self.cmd_uanalyze))
        application.add_handler(CommandHandler("subscribe", self.subscribe))
        application.add_handler(CommandHandler("unsubscribe", self.unsubscribe))
        application.add_handler(CommandHandler("news_summary", self.cmd_news_summary))

        application.add_handler(conv_handler)
        application.add_handler(CallbackQueryHandler(self.button_cb))
        # 錯誤處理
        application.add_error_handler(self.cmd_error)
        # repeat task
        application.job_queue.run_repeating(callback=self.get_reports , interval=600, first=1, data={}, name="get_reports")
        application.job_queue.run_repeating(callback=self.get_news    , interval=60,  first=1, data={}, name="get_news")
        # 註冊文字訊息處理器，這會回應用戶發送的所有文字訊息
        application.add_handler(MessageHandler(filters.ALL, self.handle_message))
        # set menu
        # asyncio.run(self.set_main_menu(application))
        self.set_main_menu(application)
        # 開始輪詢
        application.run_polling()

if __name__ == '__main__':
    MyBot = TelegramBot()
    MyBot.run()

